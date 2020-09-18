# -*- coding: utf-8 -*-
"""
Created on Mon Aug 31 00:13:49 2020

@author: Diego
"""

import glob, os
from docx import Document
from docx.shared import Inches

document = Document()
name = input("Enter name: ")
document.add_heading(name, 0)
#p = document.add_paragraph('')
#p.add_run('Diego Lopez').bold = True
document.add_heading('Diego Lopez', level=1)


path =  'C:\\Users\\Familia Lopez\\da'

os.chdir(path)

for file in os.listdir(path):
    if file.endswith(".jpeg") or file.endswith(".jpg") or file.endswith(".png")\
            or file.endswith(".JPG"):
        document.add_picture(file, width=Inches(6.5))
        print("correct!\n")

document.save(name+'.docx')



